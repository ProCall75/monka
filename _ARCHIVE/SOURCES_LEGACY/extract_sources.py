#!/usr/bin/env python3
"""
Script d'extraction des sources Monka
Extrait tous les fichiers Legacy (.docx) et Excel (.xlsx) vers JSON
"""

import os
import json
import hashlib
from datetime import datetime
from pathlib import Path

# Dépendances
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx non installé. Installer avec: pip install python-docx")

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("⚠️  openpyxl non installé. Installer avec: pip install openpyxl")

# Chemins
SOURCES_DIR = Path(__file__).parent
LEGACY_DIR = SOURCES_DIR / "legacy"
EXCEL_DIR = SOURCES_DIR / "excel"
EXTRACTED_DIR = SOURCES_DIR / "extracted"

def get_file_hash(filepath):
    """Calcule le hash MD5 d'un fichier"""
    with open(filepath, 'rb') as f:
        return hashlib.md5(f.read()).hexdigest()

def extract_docx(filepath):
    """Extrait le contenu d'un fichier Word"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx non installé"}
    
    doc = Document(filepath)
    
    # Extraire les paragraphes
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Extraire les tableaux
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables)
    }

def extract_xlsx(filepath):
    """Extrait le contenu d'un fichier Excel"""
    if not XLSX_AVAILABLE:
        return {"error": "openpyxl non installé"}
    
    wb = openpyxl.load_workbook(filepath, data_only=True)
    
    sheets = {}
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = []
        for row in sheet.iter_rows(values_only=True):
            # Convertir les valeurs en strings, gérer None
            row_data = [str(cell) if cell is not None else "" for cell in row]
            # Ne garder que les lignes avec du contenu
            if any(cell.strip() for cell in row_data):
                rows.append(row_data)
        
        if rows:
            sheets[sheet_name] = {
                "rows": rows,
                "row_count": len(rows),
                "headers": rows[0] if rows else []
            }
    
    return {
        "sheets": sheets,
        "sheet_count": len(sheets)
    }

def extract_all_sources(verify_only=False):
    """Extrait toutes les sources et génère les JSON"""
    
    manifest = {
        "extraction_date": datetime.now().isoformat(),
        "legacy_files": [],
        "excel_files": [],
        "status": "OK"
    }
    
    # Créer le dossier extracted si nécessaire
    EXTRACTED_DIR.mkdir(exist_ok=True)
    
    # === LEGACY (.docx) ===
    legacy_complete = {}
    legacy_files = list(LEGACY_DIR.glob("*.docx"))
    
    print(f"\n📄 Extraction de {len(legacy_files)} fichiers Legacy...")
    
    for docx_file in sorted(legacy_files):
        filename = docx_file.name
        file_hash = get_file_hash(docx_file)
        
        print(f"  → {filename}")
        
        if not verify_only:
            content = extract_docx(docx_file)
            legacy_complete[filename] = content
        
        manifest["legacy_files"].append({
            "filename": filename,
            "hash": file_hash,
            "status": "OK"
        })
    
    if not verify_only and legacy_complete:
        output_path = EXTRACTED_DIR / "legacy_complete.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(legacy_complete, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Généré: legacy_complete.json ({len(legacy_complete)} fichiers)")
    
    # === EXCEL (.xlsx) ===
    excel_files = list(EXCEL_DIR.glob("*.xlsx"))
    
    print(f"\n📊 Extraction de {len(excel_files)} fichiers Excel...")
    
    for xlsx_file in sorted(excel_files):
        filename = xlsx_file.name
        file_hash = get_file_hash(xlsx_file)
        
        print(f"  → {filename}")
        
        if not verify_only:
            content = extract_xlsx(xlsx_file)
            
            # Générer un JSON par fichier Excel
            safe_name = filename.replace(" ", "_").replace("(", "").replace(")", "").replace("+", "_")
            output_path = EXTRACTED_DIR / f"{safe_name}_extracted.json"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(content, f, ensure_ascii=False, indent=2)
            print(f"    ✅ Généré: {output_path.name}")
        
        manifest["excel_files"].append({
            "filename": filename,
            "hash": file_hash,
            "status": "OK"
        })
    
    # Sauvegarder le manifest
    manifest_path = EXTRACTED_DIR / "extraction_manifest.json"
    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    
    print(f"\n📋 Manifest généré: extraction_manifest.json")
    print(f"\n✅ EXTRACTION TERMINÉE")
    print(f"   Legacy: {len(manifest['legacy_files'])} fichiers")
    print(f"   Excel:  {len(manifest['excel_files'])} fichiers")
    
    return manifest

if __name__ == "__main__":
    import sys
    
    verify_only = "--verify" in sys.argv
    
    if verify_only:
        print("🔍 Mode vérification uniquement")
    else:
        print("🔄 Extraction complète des sources")
    
    extract_all_sources(verify_only=verify_only)
